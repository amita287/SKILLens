"""
Resume Parser Module
Supports PDF (pdfplumber → PyMuPDF → PyPDF2 fallback) and DOCX.
Produces clean, well-structured plain text.
"""

import io
import re


def parse_resume(uploaded_file) -> str:
    """Parse uploaded resume file (PDF or DOCX) and return raw text."""
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()

    if filename.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return _parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {filename}. Please upload a PDF or DOCX.")


# ─────────────────────────────────────────────
# PDF PARSER  (three-library cascade)
# ─────────────────────────────────────────────
def _parse_pdf(file_bytes: bytes) -> str:
    text = ""

    # 1. pdfplumber — good for tables and structured layouts
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if page_text:
                    text += page_text + "\n"
    except Exception:
        pass

    # 2. PyMuPDF (pymupdf) — best for multi-column and complex layouts
    if not text.strip():
        try:
            import pymupdf as fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                page_text = page.get_text("text")
                if page_text:
                    text += page_text + "\n"
        except Exception:
            pass

    # 3. PyPDF2 — fallback
    if not text.strip():
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception:
            pass

    if not text.strip():
        raise RuntimeError(
            "Could not extract text from PDF. "
            "The file may be scanned/image-based — please use a text-based PDF."
        )

    return _clean_text(text)


# ─────────────────────────────────────────────
# DOCX PARSER
# ─────────────────────────────────────────────
def _parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        return _clean_text("\n".join(paragraphs))

    except Exception as e:
        raise RuntimeError(f"Could not parse DOCX: {e}")


# ─────────────────────────────────────────────
# TEXT CLEANING
# ─────────────────────────────────────────────
def _clean_text(text: str) -> str:
    """Normalise extracted text without removing meaningful content."""
    # Collapse excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Normalise horizontal whitespace
    text = re.sub(r'[ \t]{2,}', ' ', text)
    # Remove non-printable characters (except newlines)
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
    # Final whitespace normalisation
    text = re.sub(r' +', ' ', text)
    return text.strip()